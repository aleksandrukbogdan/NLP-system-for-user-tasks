from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import os
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import re
import sys
from gigachat import GigaChat

# Настройка логирования с правильной кодировкой
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)

load_dotenv()

app = Flask(__name__)
CORS(app)

# Создаем директории для заявок, если их нет
REQUESTS_DIR = "requests"
DOCX_DIR = "docx_files"
UPLOADS_DIR = "uploads"
for directory in [REQUESTS_DIR, DOCX_DIR, UPLOADS_DIR]:
    if not os.path.exists(directory):
        os.makedirs(directory)

# Инициализация API ключа
GIGACHAT_API_KEY = os.getenv('GIGACHAT_API_KEY')
if not GIGACHAT_API_KEY:
    raise Exception("Не найден API ключ GigaChat. Пожалуйста, добавьте GIGACHAT_API_KEY в .env файл")

# Путь к сертификату
CA_BUNDLE_FILE = "russian_trusted_root_ca.cer"

# Проверяем наличие сертификата


# Инициализация GigaChat с сертификатом
try:
    giga = GigaChat(
        credentials=GIGACHAT_API_KEY,

        verify_ssl_certs=False
    )
    logging.info("GigaChat успешно инициализирован с сертификатом безопасности")
except Exception as e:
    logging.error(f"Ошибка инициализации GigaChat: {str(e)}")
    raise Exception(f"Не удалось инициализировать GigaChat: {str(e)}")

def validate_input(text):
    """Валидация входного текста"""
    if not text or not isinstance(text, str):
        return False, "Пустой или некорректный текст"
    if len(text.strip()) < 10:
        return False, "Текст слишком короткий"
    return True, ""

def preprocess_text(text):
    """Предварительная обработка текста"""
    # Удаление лишних пробелов
    text = re.sub(r'\s+', ' ', text).strip()
    # Удаление специальных символов
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    return text

def generate_chat_response(messages):
    try:
        # Последнее сообщение от пользователя
        user_prompt = messages[-1]['text']

        # Формируем системный промпт
        system_prompt = """# Ты - ИИ-ассистент для сбора требований на разработку ПО.
        Твоя задача — вести диалог с пользователем, чтобы собрать всю необходимую информацию для составления технического задания (ТЗ).

        ## Инструкция
        1.  Начни с приветствия и попроси пользователя описать свою задачу.
        2.  После того как пользователь опишет задачу, классифицируй её, но не сообщай категорию пользователю.
        3.  Задавай **по одному** уточняющему вопросу, чтобы собрать больше деталей. Не задавай все вопросы сразу.
        4.  Если для понимания задачи могут быть полезны какие-либо файлы (например, макеты, диаграммы, текущие документы), попроси пользователя прикрепить их.
        5.  Твои вопросы должны быть четкими и направленными на сбор информации для ТЗ (цели, задачи, функционал, сроки, бюджет и т.д.).
        6.  Веди диалог естественно. Когда получишь ответ, поблагодари и задай следующий вопрос.
        7.  **ЗАПРЕЩЕНО** использовать в ответах любые служебные слова в квадратных скобках, такие как `[CATEGORY]` или `[question]`. Общайся как человек.
        8.  Когда соберешь достаточно информации, заверши диалог. Твой финальный ответ ДОЛЖЕН начинаться со специальной фразы: `[DOCUMENT_READY]`. Этой фразы пользователь не увидит. Пример твоего финального сообщения: `[DOCUMENT_READY]Спасибо за ответы! Я собрал достаточно информации и подготовил документ. Его можно скачать по ссылке:`

        ## Формат ответа
        Твой ответ должен быть только текстом следующего сообщения в чате.
        """
        
        # Преобразуем историю сообщений в формат, понятный для GigaChat
        history = []
        for msg in messages:
            if msg['sender'] == 'user':
                history.append({'role': 'user', 'content': msg['text']})
            else:
                history.append({'role': 'assistant', 'content': msg['text']})
        
        # Собираем payload для GigaChat
        payload = {
            "model": "GigaChat",
            "messages": [
                {"role": "system", "content": system_prompt},
                *history
            ]
        }
        
        logging.info("Отправка запроса к GigaChat")
        response = giga.chat(payload)
        
        generated_text = response.choices[0].message.content
        return generated_text.strip()
        
    except Exception as e:
        logging.error(f"Ошибка генерации текста: {str(e)}")
        raise Exception(f"Ошибка при генерации текста: {str(e)}")

def create_docx_from_chat(filename, messages):
    try:
        doc = Document()
        
        # Настройка стилей
        title_style = doc.styles['Title']
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        
        heading_style = doc.styles['Heading 1']
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Добавление заголовка
        title = doc.add_heading('Техническое задание', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавление даты
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f'Дата создания: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        doc.add_heading('История обсуждения:', level=1)
        
        full_chat_history = ""
        for msg in messages:
            sender = "Пользователь" if msg['sender'] == 'user' else "Ассистент"
            full_chat_history += f"{sender}: {msg['text']}\n"
        
        doc.add_paragraph(full_chat_history)
        
        # TODO: Добавить логику для извлечения структурированной информации из чата
        # Например, можно попросить GigaChat сделать саммари в формате JSON

        doc.save(filename)
        logging.info(f"Документ сохранен: {filename}")
        return True
    except Exception as e:
        logging.error(f"Ошибка создания DOCX из чата: {str(e)}")
        return False

@app.route('/api/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "Файл не найден"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"success": False, "error": "Файл не выбран"}), 400
        
        if file:
            # Используем werkzeug для безопасного имени файла
            from werkzeug.utils import secure_filename
            filename = secure_filename(file.filename)
            file.save(os.path.join(UPLOADS_DIR, filename))
            logging.info(f"Файл '{filename}' успешно загружен.")
            return jsonify({"success": True, "filename": filename})
            
    except Exception as e:
        logging.error(f"Ошибка загрузки файла: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/chat', methods=['POST'])
def chat():
    try:
        logging.info("Получен POST запрос в /api/chat")
        data = request.json
        
        if not data or 'messages' not in data:
            return jsonify({"success": False, "error": "Отсутствуют сообщения в запросе"}), 400
            
        messages = data.get('messages')
        
        # Валидация последнего сообщения
        last_message = messages[-1]['text']
        is_valid, error_message = validate_input(last_message)
        if not is_valid:
            return jsonify({"success": False, "error": error_message}), 400
            
        # Генерация ответа от чат-бота
        try:
            bot_response = generate_chat_response(messages)
            
            # Проверяем, готов ли документ
            if bot_response.startswith('[DOCUMENT_READY]'):
                bot_response = bot_response.replace('[DOCUMENT_READY]', '').strip()
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                docx_filename = f"ТЗ_{timestamp}.docx"
                docx_filepath = os.path.join(DOCX_DIR, docx_filename)

                if create_docx_from_chat(docx_filepath, messages):
                    # Добавляем ссылку на скачивание в ответ
                    download_url = f"/downloads/{docx_filename}"
                    final_response = f"{bot_response}\n[Скачать документ]({download_url})"
                    return jsonify({"success": True, "reply": final_response, "document_ready": True})
                else:
                    return jsonify({"success": False, "error": "Не удалось создать документ"}), 500

            return jsonify({"success": True, "reply": bot_response})
        except Exception as e:
            logging.error(f"Ошибка при обработке запроса в чате: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500
    
    except Exception as e:
        print(f"Общая ошибка в /api/chat: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/downloads/<filename>')
def download_file(filename):
    return send_from_directory(DOCX_DIR, filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True) 